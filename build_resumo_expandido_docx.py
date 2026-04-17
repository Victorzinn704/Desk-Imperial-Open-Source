"""Gera versão .docx editável do Resumo Expandido.

Mesmo conteúdo do PDF — útil para preencher orientador/instituição/e-mail
antes da exportação final.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

import build_resumo_expandido as src

OUTPUT = r"c:\Users\Desktop\Documents\desk-imperial\resumo-expandido-desk-imperial.docx"


def set_font(run, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold


def add_para(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False,
             first_line_indent=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = alignment
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    # strip simple <b>...</b> tags
    segments = []
    i = 0
    while i < len(text):
        if text[i:i+3] == "<b>":
            j = text.find("</b>", i)
            segments.append((text[i+3:j], True))
            i = j + 4
        else:
            j = text.find("<b>", i)
            if j == -1:
                segments.append((text[i:], False))
                break
            segments.append((text[i:j], False))
            i = j
    for seg_text, seg_bold in segments:
        run = p.add_run(seg_text.replace("<sup>", "").replace("</sup>", ""))
        set_font(run, bold=(bold or seg_bold))
    return p


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    add_para(doc, src.TITULO, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_para(doc, src.CATEGORIA, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, src.AUTOR, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, src.ORIENTADOR, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, src.FILIACAO, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, src.EMAIL, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    add_para(doc, src.RESUMO)
    add_para(doc, src.PALAVRAS_CHAVE)

    add_para(doc, "1. INTRODUÇÃO", alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=True, space_after=6)
    for par in src.INTRO:
        add_para(doc, par, first_line_indent=Cm(1.25))

    add_para(doc, "2. MÉTODO", alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=True, space_after=6)
    for par in src.METODO:
        add_para(doc, par, first_line_indent=Cm(1.25))

    add_para(doc, "3. RESULTADOS E DISCUSSÃO", alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=True, space_after=6)
    for par in src.RESULTADOS:
        add_para(doc, par, first_line_indent=Cm(1.25))

    add_para(doc, "4. CONSIDERAÇÕES FINAIS", alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=True, space_after=6)
    for par in src.CONSIDERACOES:
        add_para(doc, par, first_line_indent=Cm(1.25))

    add_para(doc, "REFERÊNCIAS", alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=True, space_after=6)
    for ref in src.REFERENCIAS:
        add_para(doc, ref)

    doc.save(OUTPUT)
    print("OK:", OUTPUT)


if __name__ == "__main__":
    build()
